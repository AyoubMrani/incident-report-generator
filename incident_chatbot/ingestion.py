import json
import os
import re
from typing import Callable

import numpy as np
import streamlit as st
from docx import Document

from .config import REPORTS_FOLDER, EMBED_MODEL_NAME


def _read_docx(path: str) -> str:
    doc = Document(path)
    parts = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


def _read_json(path: str) -> str:
    with open(path, "r", encoding="utf-8") as f:
        data = json.load(f)

    def _flatten(obj, prefix=""):
        lines = []
        if isinstance(obj, dict):
            for key, value in obj.items():
                lines.extend(_flatten(value, f"{prefix}{key}: "))
        elif isinstance(obj, list):
            for item in obj:
                lines.extend(_flatten(item, prefix))
        else:
            lines.append(f"{prefix}{obj}")
        return lines

    return "\n".join(_flatten(data))


def _read_md(path: str) -> str:
    with open(path, "r", encoding="utf-8") as f:
        return f.read()


def _chunk(text: str, size: int = 700, overlap: int = 120) -> list[str]:
    text = re.sub(r"\s+", " ", text).strip()
    chunks = []
    start = 0
    while start < len(text):
        chunks.append(text[start : start + size])
        start += size - overlap
    return chunks


def _get_report_title(path: str, content: str) -> str:
    filename = os.path.splitext(os.path.basename(path))[0]
    if path.endswith(".json"):
        try:
            data = json.loads(content)
            metadata = data.get("metadata", {})
            title = metadata.get("title") or metadata.get("incident_id")
            if title:
                return str(title)
        except Exception:
            pass
    if path.endswith(".md"):
        for line in content.splitlines():
            if line.startswith("#"):
                return line.lstrip("#").strip()
    match = re.search(r"INC\d+", content, re.IGNORECASE)
    if match:
        return match.group(0).upper()
    return filename


def _extract_incident_id_from_text(text: str | None) -> str | None:
    if not text:
        return None
    match = re.search(r"INC\d+", str(text), re.IGNORECASE)
    return match.group(0).upper() if match else None


@st.cache_resource(show_spinner="Building knowledge base…")
def build_knowledge_base():
    from sentence_transformers import SentenceTransformer

    loaders: dict[str, Callable[[str], str]] = {
        ".docx": _read_docx,
        ".json": _read_json,
        ".md": _read_md,
        ".txt": _read_md,
    }

    documents: list[str] = []
    metadata: list[dict] = []
    roots = [REPORTS_FOLDER] if os.path.isdir(REPORTS_FOLDER) else []

    if not roots:
        raise ValueError(
            f"No docs/ or reports/ folder found at {os.path.abspath('.')}. "
            "Create at least one and add incident documents."
        )

    n_files = 0
    for root in roots:
        for filename in sorted(os.listdir(root)):
            ext = os.path.splitext(filename)[1].lower()
            if ext not in loaders:
                continue
            path = os.path.join(root, filename)
            source = f"{os.path.basename(root)}/{filename}"
            try:
                content = loaders[ext](path)
                title = _get_report_title(path, content)
                incident_id = _extract_incident_id_from_text(content)
                for chunk_id, chunk in enumerate(_chunk(content)):
                    documents.append(chunk)
                    metadata.append({
                        "source": source,
                        "path": path,
                        "title": title,
                        "chunk_id": chunk_id,
                        "incident_id": incident_id,
                    })
                n_files += 1
            except Exception as exc:
                st.warning(f"Could not index {source}: {exc}")

    if not documents:
        raise ValueError("Knowledge base is empty — add .docx / .json / .md files to docs/ or reports/.")

    model = SentenceTransformer(EMBED_MODEL_NAME)
    embeddings = model.encode(documents, convert_to_numpy=True, show_progress_bar=False).astype("float32")
    return model, embeddings, documents, metadata, n_files
