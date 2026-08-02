"""
chatbot/ingestion.py — build the retrieval knowledge base from report files.

Ported from the original incident_chatbot/ingestion.py with two changes needed
to run headless under FastAPI:

  1. Streamlit removed. The `@st.cache_resource` decorator on
     build_knowledge_base is gone — caching now happens by building the KB once
     in the app lifespan and holding it on app.state (see main.py). `st.warning`
     on an unindexable file becomes a collected warning on the returned object.

  2. The reports directory is a parameter, not the hard-coded REPORTS_FOLDER, so
     the chatbot indexes the same shared reports/ dir the generator writes to.

The file readers, chunking, title extraction, and INC-id extraction are
unchanged.
"""

from __future__ import annotations

import hashlib
import json
import os
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Callable

import numpy as np

from .config import CHUNK_OVERLAP, CHUNK_SIZE, EMBED_MODEL_NAME


@dataclass
class KnowledgeBase:
    """Everything retrieval.search() needs, built once and reused per process."""

    embed_model: object
    embeddings: np.ndarray
    documents: list[str]
    metadata: list[dict]
    n_files: int
    bm25: object | None = None   # BM25Index over `documents`; fused with vectors
    warnings: list[str] = field(default_factory=list)


# ── File readers (unchanged) ──────────────────────────────────────────────────


def _read_docx(path: str) -> str:
    # Lazy import: python-docx is only needed when a .docx report is indexed.
    from docx import Document

    doc = Document(path)
    parts = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


def _strip_html(text: str) -> str:
    """Quill stores paragraph content as HTML; reduce it to readable text.

    Inline <img> tags (screenshots pasted into a step-by-step guide, often as
    base64 data URIs) are replaced with a positional marker rather than dropped.
    The image bytes are useless to a text model and would swamp the chunk, but
    the *fact* that a screenshot documents this point in the procedure matters —
    it lets the answer acknowledge the visual instead of pretending it is absent.
    """
    raw = text or ""
    # Preserve list-item boundaries that would otherwise collapse into one line.
    raw = re.sub(r"</li\s*>", " \n", raw, flags=re.IGNORECASE)
    raw = re.sub(r"<br\s*/?>", " \n", raw, flags=re.IGNORECASE)

    # Mark images in place, numbered in document order.
    counter = {"n": 0}

    def _img_marker(_match: re.Match) -> str:
        counter["n"] += 1
        return f" [SCREENSHOT {counter['n']}] "

    raw = re.sub(r"<img\b[^>]*>", _img_marker, raw, flags=re.IGNORECASE)

    text = re.sub(r"<[^>]+>", " ", raw)
    text = (
        text.replace("&amp;", "&").replace("&lt;", "<").replace("&gt;", ">")
        .replace("&quot;", '"').replace("&#39;", "'").replace("&nbsp;", " ")
    )
    # Collapse runs of spaces/tabs but keep the line breaks we inserted.
    text = re.sub(r"[ \t]+", " ", text)
    return re.sub(r"\n\s*\n+", "\n", text).strip()


def _render_block(block: dict) -> str:
    """Turn one report block into clean, human-readable text.

    This is the accuracy fix: the previous flattener emitted structural noise
    ('id: b4', 'type: heading', 'level: 2', base64 image blobs) that polluted
    embeddings and the LLM context. Here we extract only the incident *content*
    per block type, so retrieval matches on meaning and the model sees real text.
    """
    btype = block.get("type")
    title = (block.get("title") or "").strip()
    out: list[str] = []

    if btype == "heading":
        out.append(f"## {block.get('content', '').strip()}")
    elif btype == "paragraph":
        if title:
            out.append(f"{title}:")
        out.append(_strip_html(block.get("content", "")))
    elif btype == "list":
        if title:
            out.append(f"{title}:")
        if block.get("label"):
            out.append(str(block["label"]).strip())
        for item in block.get("items", []):
            out.append(f"- {_strip_html(str(item))}")
    elif btype == "code":
        # A code block may hold snippets and descriptions (see report schema).
        if title:
            out.append(f"{title}:")
        for item in block.get("items", []) if isinstance(block.get("items"), list) else []:
            if isinstance(item, dict):
                if item.get("header"):
                    out.append(str(item["header"]).strip())
                if item.get("content"):
                    out.append(str(item["content"]).strip())
        # Some code blocks store content directly.
        if block.get("content"):
            out.append(str(block["content"]).strip())
    elif btype == "table":
        if title:
            out.append(f"{title}:")
        headers = block.get("headers", [])
        if headers:
            out.append(" | ".join(str(h) for h in headers))
        for row in block.get("rows", []):
            out.append(" | ".join(str(c) for c in row))
    elif btype in ("incident_example",):
        ref = block.get("incident_id", "")
        link = block.get("link", "")
        out.append(f"Related incident: {ref} {link}".strip())
    elif btype in ("description_box",):
        if title:
            out.append(f"{title}:")
        out.append(_strip_html(block.get("content", "")))
    elif btype == "image":
        # Skip the base64 data_url; keep only a human caption if present.
        cap = (block.get("caption") or "").strip()
        if cap:
            out.append(f"[image: {cap}]")

    return "\n".join(p for p in out if p)


def _read_json(path: str) -> str:
    """Extract clean incident text from a report JSON (schema-aware).

    Handles both on-disk shapes: flat {metadata, blocks} and the legacy wrapper
    {report: {metadata, blocks}}. Emits title, key metadata, and the readable
    content of each block — no structural keys, no base64.
    """
    with open(path, "r", encoding="utf-8") as f:
        data = json.load(f)

    # Unwrap the legacy {editingFilename, markdown, report:{...}} shape.
    report = data.get("report") if isinstance(data.get("report"), dict) else data
    metadata = report.get("metadata", {}) if isinstance(report, dict) else {}
    blocks = report.get("blocks", []) if isinstance(report, dict) else []

    lines: list[str] = []
    # Lead with the identifying metadata so it's searchable and grounds the model.
    if metadata.get("title"):
        lines.append(f"# {metadata['title']}")
    for key in ("incident_id", "category", "subcategory", "caller", "date", "priority"):
        if metadata.get(key):
            lines.append(f"{key}: {metadata[key]}")

    for block in blocks:
        if isinstance(block, dict):
            rendered = _render_block(block)
            if rendered:
                lines.append(rendered)

    # Fallback: if this wasn't a report-shaped JSON at all, keep the markdown.
    if not blocks and isinstance(data.get("markdown"), str):
        lines.append(data["markdown"])

    return "\n".join(lines)


def _read_md(path: str) -> str:
    with open(path, "r", encoding="utf-8") as f:
        return f.read()


def _chunk(text: str, size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> list[str]:
    text = re.sub(r"\s+", " ", text).strip()
    chunks = []
    start = 0
    while start < len(text):
        chunks.append(text[start : start + size])
        start += size - overlap
    return chunks


def _get_report_title(path: str, content: str) -> str:
    filename = os.path.splitext(os.path.basename(path))[0]
    if path.endswith(".json"):
        # Read the file, not `content`: by this point content is the extracted
        # prose from _read_json, so json.loads() on it always fails and every
        # report would fall through to the bare "INC…" id below — losing the
        # human title that source selection scores entity overlap against.
        try:
            with open(path, "r", encoding="utf-8") as f:
                data = json.load(f)
            report = data.get("report") if isinstance(data.get("report"), dict) else data
            metadata = (report or {}).get("metadata", {}) or {}
            title = metadata.get("title")
            if title and str(title).strip():
                return str(title).strip()
        except Exception:
            metadata = {}
        # No declared title. The filename usually carries a real one
        # ("INC1048202_Yellow Duplicate Cleanup") — prefer it over a bare id,
        # since source selection scores entity overlap against the title.
        stem = re.sub(r"^INC\d+[_\-\s]*", "", filename, flags=re.IGNORECASE).strip()
        if stem and not stem.lower().startswith("incident_untitled"):
            return stem.replace("_", " ").strip()
        declared = (metadata or {}).get("incident_id")
        if declared and str(declared).strip():
            return str(declared).strip()
    if path.endswith(".md"):
        for line in content.splitlines():
            if line.startswith("#"):
                return line.lstrip("#").strip()
    match = re.search(r"INC\d+", content, re.IGNORECASE)
    if match:
        return match.group(0).upper()
    return filename


def _extract_incident_id_from_text(text: str | None) -> str | None:
    if not text:
        return None
    match = re.search(r"INC\d+", str(text), re.IGNORECASE)
    return match.group(0).upper() if match else None


def _declared_incident_id(path: str) -> str | None:
    """The id a report declares in its own metadata, if it is a report JSON.

    Preferred over scanning the text: a report may *reference* other incidents
    (an `incident_example` block links a related ticket), and a raw text scan
    returns whichever id appears first — so a report would get cited under its
    neighbour's number. Metadata is the authoritative self-identifier.
    """
    if not path.lower().endswith(".json"):
        return None
    try:
        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f)
    except Exception:  # noqa: BLE001 — fall back to the text scan
        return None
    report = data.get("report") if isinstance(data.get("report"), dict) else data
    if not isinstance(report, dict):
        return None
    declared = (report.get("metadata") or {}).get("incident_id")
    declared = str(declared).strip() if declared else ""
    return declared or None


# ── Knowledge base builder (Streamlit removed) ────────────────────────────────


def _cache_dir() -> Path:
    return Path(os.getenv("EMBED_CACHE_DIR", "data/embed-cache"))


def _embedding_cache_key(documents: list[str]) -> str:
    """Digest identifying this exact corpus under this exact embedding model."""
    h = hashlib.sha256()
    h.update(EMBED_MODEL_NAME.encode())
    for doc in documents:
        h.update(b"\0")
        h.update(doc.encode("utf-8", "replace"))
    return h.hexdigest()[:32]


def _load_cached_embeddings(key: str, n_documents: int):
    """Cached vectors for this corpus, or None when absent/stale/unreadable."""
    path = _cache_dir() / f"{key}.npy"
    if not path.is_file():
        return None
    try:
        cached = np.load(path)
    except Exception:  # noqa: BLE001 — a corrupt cache must never break startup
        return None
    # Guard against a truncated write: shape must still match the corpus.
    if cached.ndim != 2 or cached.shape[0] != n_documents:
        return None
    return cached.astype("float32")


def _store_cached_embeddings(key: str, embeddings) -> None:
    """Persist vectors for the next boot; failures are non-fatal by design."""
    try:
        cache = _cache_dir()
        cache.mkdir(parents=True, exist_ok=True)
        # Write to a temp file then rename so a crash can't leave a half file
        # that a later boot would read as valid.
        tmp = cache / f".{key}.tmp.npy"
        np.save(tmp, embeddings)
        tmp.replace(cache / f"{key}.npy")
        for stale in cache.glob("*.npy"):
            if stale.name != f"{key}.npy":
                stale.unlink(missing_ok=True)
    except Exception:  # noqa: BLE001 — caching is an optimisation, not a contract
        pass


def build_knowledge_base(reports_dir: str | Path) -> KnowledgeBase:
    """Index every supported file under ``reports_dir`` into a KnowledgeBase.

    Raises ValueError if the directory is missing or yields no indexable
    documents (same failure conditions as the original).
    """
    from sentence_transformers import SentenceTransformer

    reports_dir = str(reports_dir)

    loaders: dict[str, Callable[[str], str]] = {
        ".docx": _read_docx,
        ".json": _read_json,
        ".md": _read_md,
        ".txt": _read_md,
    }

    if not os.path.isdir(reports_dir):
        raise ValueError(
            f"Reports directory not found: {os.path.abspath(reports_dir)}. "
            "Create it and add incident documents."
        )

    documents: list[str] = []
    metadata: list[dict] = []
    warnings: list[str] = []
    n_files = 0

    root_label = os.path.basename(os.path.normpath(reports_dir))
    entries = sorted(os.listdir(reports_dir))
    # The generator writes every report twice: the structured .json and a
    # flattened .md rendering of the same incident. Indexing both stores each
    # incident twice, so near-identical duplicates compete for the handful of
    # source slots the answer gets — and the .md loses the schema (block types,
    # screenshot markers, declared incident id) that makes the .json accurate.
    # The .json is authoritative; skip a .md that merely mirrors one.
    json_stems = {
        os.path.splitext(name)[0]
        for name in entries
        if name.lower().endswith(".json")
    }

    for filename in entries:
        ext = os.path.splitext(filename)[1].lower()
        if ext not in loaders:
            continue
        if ext == ".md" and os.path.splitext(filename)[0] in json_stems:
            continue
        path = os.path.join(reports_dir, filename)
        source = f"{root_label}/{filename}"
        try:
            content = loaders[ext](path)
            title = _get_report_title(path, content)
            # Priority: the report's own metadata, then the filename (reports are
            # saved as "<INC id>_<title>.<ext>" pairs, so the .md twin of a JSON
            # report still identifies itself correctly), and only then a text
            # scan — which can pick up a *referenced* incident instead.
            incident_id = (
                _declared_incident_id(path)
                or _extract_incident_id_from_text(filename)
                or _extract_incident_id_from_text(content)
            )
            for chunk_id, chunk in enumerate(_chunk(content)):
                documents.append(chunk)
                metadata.append(
                    {
                        "source": source,
                        "path": path,
                        "title": title,
                        "chunk_id": chunk_id,
                        "incident_id": incident_id,
                    }
                )
            n_files += 1
        except Exception as exc:  # noqa: BLE001 — keep indexing other files
            warnings.append(f"Could not index {source}: {exc}")

    if not documents:
        raise ValueError(
            "Knowledge base is empty — add .docx / .json / .md files to the "
            f"reports directory ({os.path.abspath(reports_dir)})."
        )

    model = SentenceTransformer(EMBED_MODEL_NAME)
    # Encoding 200 chunks costs ~1s on every boot even though the corpus rarely
    # changes. Cache the vectors under a digest of (model, chunk texts) so an
    # unchanged reports/ dir reloads instantly and any edit invalidates the key.
    cache_key = _embedding_cache_key(documents)
    embeddings = _load_cached_embeddings(cache_key, len(documents))
    if embeddings is None:
        embeddings = (
            model.encode(documents, convert_to_numpy=True, show_progress_bar=False)
            .astype("float32")
        )
        _store_cached_embeddings(cache_key, embeddings)
    # Build the lexical BM25 index over the same chunks so retrieval can fuse
    # exact-term matching (INC ids, table/function names) with semantic search.
    from .bm25 import BM25Index

    bm25 = BM25Index(documents)

    return KnowledgeBase(
        embed_model=model,
        embeddings=embeddings,
        documents=documents,
        metadata=metadata,
        n_files=n_files,
        bm25=bm25,
        warnings=warnings,
    )
