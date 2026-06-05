"""
ServiceNow Incident Analysis Chatbot
=====================================
Architecture:
  Layer 1 — Ingestion:   DOCX + JSON + MD → chunks → embeddings → vector store
  Layer 2 — Extraction:  VLM (screenshot) or text parser → structured incident JSON
  Layer 3 — Retrieval:   query vector store → top-k relevant chunks
  Layer 4 — Resolution:  incident JSON + chunks → deterministic steps + Ollama narrative
  Layer 5 — Output:      structured JSON + readable UI + log SQL

Tabs:
    • Main Demo          — extraction → retrieval → resolution, all in one
  • Knowledge Base     — inspect what's indexed
  • Qwen2-VL Test      — raw vision model sandbox
"""

import json
import os
import re
import tempfile

import numpy as np
import streamlit as st
from docx import Document
from PIL import Image, ImageDraw

# ──────────────────────────────────────────────────────────────────────────────
# CONFIG
# ──────────────────────────────────────────────────────────────────────────────
DOCS_FOLDER = "docs"
REPORTS_FOLDER = "reports"
TOP_K = 4
OLLAMA_MODEL = "llama3:8b"
QWEN_MLX_MODEL_PATH = os.path.join("models", "qwen2.5-vl-7b-4bit-vlm")
QWEN_HF_MODEL_ID = "Qwen/Qwen2.5-VL-7B-Instruct"
EMBED_MODEL_NAME = "all-MiniLM-L6-v2"

os.environ.setdefault("TOKENIZERS_PARALLELISM", "false")

HOME_ID_PATTERN = re.compile(r"\b[A-Z]{3}\d{12}\b")

# ──────────────────────────────────────────────────────────────────────────────
# LAYER 1 — INGESTION
# ──────────────────────────────────────────────────────────────────────────────

def _extract_docx(path: str) -> str:
    doc = Document(path)
    parts = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


def _extract_json(path: str) -> str:
    with open(path, "r", encoding="utf-8") as f:
        data = json.load(f)
    # Flatten JSON to readable text so it embeds meaningfully
    def _flatten(obj, prefix=""):
        lines = []
        if isinstance(obj, dict):
            for k, v in obj.items():
                lines.extend(_flatten(v, f"{prefix}{k}: "))
        elif isinstance(obj, list):
            for item in obj:
                lines.extend(_flatten(item, prefix))
        else:
            lines.append(f"{prefix}{obj}")
        return lines
    return "\n".join(_flatten(data))


def _extract_md(path: str) -> str:
    with open(path, "r", encoding="utf-8") as f:
        return f.read()


def _chunk(text: str, size: int = 700, overlap: int = 120) -> list[str]:
    text = re.sub(r"\s+", " ", text).strip()
    chunks, start = [], 0
    while start < len(text):
        chunks.append(text[start : start + size])
        start += size - overlap
    return chunks


def _extract_incident_id_from_text(text: str | None) -> str | None:
    if not text:
        return None
    match = re.search(r"INC\d+", str(text), re.IGNORECASE)
    if match:
        return match.group(0).upper()
    return None


def _extract_incident_id_from_report(path: str, content: str | None = None) -> str | None:
    try:
        if path.lower().endswith(".json"):
            with open(path, "r", encoding="utf-8") as f:
                data = json.load(f)
            metadata = data.get("metadata", {}) if isinstance(data, dict) else {}
            if metadata.get("incident_id"):
                return _normalize_incident_id(metadata.get("incident_id"))
        if content is None:
            with open(path, "r", encoding="utf-8") as f:
                content = f.read()
        return _extract_incident_id_from_text(content)
    except Exception:
        return None


@st.cache_resource(show_spinner="Building knowledge base…")
def build_knowledge_base():
    """
    Index all supported files in docs/ and reports/.
    Returns: (embed_model, embeddings_matrix, documents, metadata)
    """
    from sentence_transformers import SentenceTransformer

    documents, metadata = [], []

    source_roots = [root for root in [DOCS_FOLDER, REPORTS_FOLDER] if os.path.isdir(root)]
    if not source_roots:
        raise ValueError(
            f"No knowledge folders found. Expected at least one of: {os.path.abspath(DOCS_FOLDER)}, {os.path.abspath(REPORTS_FOLDER)}"
        )

    loaders = {
        ".docx": _extract_docx,
        ".json": _extract_json,
        ".md":   _extract_md,
        ".txt":  _extract_md,
    }

    files_indexed = 0
    for root in source_roots:
        for filename in sorted(os.listdir(root)):
            ext = os.path.splitext(filename)[1].lower()
            if ext not in loaders:
                continue
            path = os.path.join(root, filename)
            source_label = f"{os.path.basename(root)}/{filename}"
            try:
                content = loaders[ext](path)
                incident_id = _extract_incident_id_from_report(path, content)
                for i, chunk in enumerate(_chunk(content)):
                    documents.append(chunk)
                    metadata.append({"source": source_label, "chunk_id": i, "ext": ext, "root": root, "incident_id": incident_id})
                files_indexed += 1
            except Exception as e:
                st.warning(f"Could not index {source_label}: {e}")

    if not documents:
        raise ValueError("No documents found in docs/ or reports/. Add .docx, .json, .md, or .txt files.")

    model = SentenceTransformer(EMBED_MODEL_NAME)
    embeddings = model.encode(documents, convert_to_numpy=True, show_progress_bar=False).astype("float32")

    return model, embeddings, documents, metadata, files_indexed


# ──────────────────────────────────────────────────────────────────────────────
# LAYER 2 — EXTRACTION (text path; VLM path handled inline in UI)
# ──────────────────────────────────────────────────────────────────────────────

# Extraction prompt used for BOTH VLM (screenshot) and Ollama (text)
EXTRACTION_SYSTEM_PROMPT = """You are an IT incident data extractor for an NRI operations team.
Return ONLY valid JSON — no markdown, no preamble, no explanation outside the JSON.

Schema:
{
  "incident_id": "INC number or null",
  "caller": "full name or null",
  "category": "string or null",
  "subcategory": "string or null",
  "service": "string or null",
  "priority": "string or null",
  "incident_type": "DELETE_PROVISION | DELETE_DEFECTIVE_PORT | DELETE_DUPLICATE_RECORDS | CHANGE_HOME_STATUS | UNKNOWN — or list if multiple",
  "short_description": "string or null",
  "full_description": "string or null",
  "home_ids": ["list of NRI home IDs — pattern 3 uppercase letters + 12 digits"],
  "access_numbers": ["list of numeric access numbers"],
  "target_table": "fm_opv | of_networkCreation | null",
  "action_required": "plain English summary",
  "status_to_delete": "P | D | P and D | null",
  "target_status": "home status to change TO or null",
  "source_status": "home status to change FROM or null",
  "confidence": "HIGH | MEDIUM | LOW"
}

Classification rules:
- Delete P from status → DELETE_PROVISION
- Delete D from status → DELETE_DEFECTIVE_PORT  
- Both P and D → ["DELETE_PROVISION","DELETE_DEFECTIVE_PORT"]
- Duplicate rows / duplicate cleanup workflow → DELETE_DUPLICATE_RECORDS
- Home status change / planning state change → CHANGE_HOME_STATUS
- None match → UNKNOWN

For screenshots: extract every visible field verbatim. Never hallucinate values not visible in the image."""

VLM_EXTRACTION_PROMPT = """This is a screenshot of a ServiceNow incident form or a database/spreadsheet UI.

Step 1 — List every visible field label and its value exactly as shown (write 'not visible' if unclear).
Step 2 — Return the extraction JSON below. Do not describe people or scenes.

""" + EXTRACTION_SYSTEM_PROMPT


def extract_from_text_via_ollama(text: str) -> dict:
    """Use Ollama to extract structured JSON from plain text incident."""
    prompt = f"{EXTRACTION_SYSTEM_PROMPT}\n\nIncident text:\n{text}"
    raw = ask_ollama(prompt)
    return _parse_json_safe(raw)


def _parse_json_safe(raw: str) -> dict:
    """Strip markdown fences and parse JSON, returning error dict on failure."""
    cleaned = re.sub(r"```(?:json)?", "", raw).strip().rstrip("`").strip()
    # Find first { ... } block
    match = re.search(r"\{.*\}", cleaned, re.DOTALL)
    if match:
        cleaned = match.group(0)
    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        return {"error": "Could not parse JSON from model output", "raw": raw[:500]}


def _normalize_incident_id(value) -> str | None:
    if not value:
        return None
    match = re.search(r"INC\d+", str(value), re.IGNORECASE)
    if match:
        return match.group(0).upper()
    text = str(value).strip().upper()
    return text or None


def _find_exact_report_path(incident_id: str | None) -> str | None:
    incident_id = _normalize_incident_id(incident_id)
    if not incident_id or not os.path.isdir(REPORTS_FOLDER):
        return None
    for filename in sorted(os.listdir(REPORTS_FOLDER)):
        if not filename.lower().endswith((".json", ".md", ".txt")):
            continue
        path = os.path.join(REPORTS_FOLDER, filename)
        if _normalize_incident_id(_extract_incident_id_from_report(path)) == incident_id:
            return path
    return None


def _summarize_exact_report(report_data: dict) -> dict:
    metadata = report_data.get("metadata", {}) if isinstance(report_data, dict) else {}
    blocks = report_data.get("blocks", []) if isinstance(report_data, dict) else []

    text_parts: list[str] = []
    home_ids: list[str] = []
    access_numbers: list[str] = []
    action_lines: list[str] = []
    incident_type = "UNKNOWN"
    target_table = None
    status_to_delete = None
    target_status = None
    source_status = None
    code_blocks: list[dict] = []

    for block in blocks:
        block_type = block.get("type")
        if block_type in {"heading", "paragraph"}:
            content = str(block.get("content", "")).strip()
            if content:
                text_parts.append(content)
        elif block_type == "list":
            items = [str(item).strip() for item in block.get("items", []) if str(item).strip()]
            text_parts.extend(items)
            action_lines.extend(items)
        elif block_type == "code":
            code = str(block.get("content", "")).strip()
            if code:
                text_parts.append(code)
                code_blocks.append({"content": code, "language": block.get("language")})
        elif block_type == "table":
            rows = block.get("rows", []) or []
            headers = [str(header).strip().lower() for header in block.get("headers", []) or []]
            for row in rows:
                if not isinstance(row, list):
                    continue
                if row:
                    first = str(row[0]).strip()
                    if HOME_ID_PATTERN.search(first) and first not in home_ids:
                        home_ids.append(first)
                if len(row) > 1:
                    second = str(row[1]).strip()
                    if second.isdigit() and second not in access_numbers:
                        access_numbers.append(second)
            if any("home id" in header for header in headers):
                target_table = target_table or "fm_opv"
        else:
            raw_value = block.get("content")
            if raw_value:
                text_parts.append(str(raw_value))

    report_text = "\n".join(text_parts)
    report_text_lower = report_text.lower()
    title = metadata.get("title") or metadata.get("incident_id")

    if incident_type == "UNKNOWN":
        if "duplicate" in report_text_lower:
            incident_type = "DELETE_DUPLICATE_RECORDS"
        elif "defective port" in report_text_lower:
            incident_type = "DELETE_DEFECTIVE_PORT"
        elif "home status" in report_text_lower or "change status" in report_text_lower or "planning" in report_text_lower:
            incident_type = "CHANGE_HOME_STATUS"
        elif "provision" in report_text_lower or "status" in report_text_lower:
            incident_type = "DELETE_PROVISION"

    if target_table is None:
        if any("fm_opv" in block.get("content", "").lower() for block in blocks if block.get("type") == "code"):
            target_table = "fm_opv"
        elif any("of_networkcreation" in block.get("content", "").lower() for block in blocks if block.get("type") == "code"):
            target_table = "of_networkCreation"

    return {
        "incident_id": metadata.get("incident_id"),
        "caller": metadata.get("caller"),
        "category": metadata.get("category"),
        "subcategory": metadata.get("subcategory"),
        "service": metadata.get("service"),
        "priority": metadata.get("priority"),
        "incident_type": incident_type,
        "short_description": title,
        "full_description": report_text[:3000] or None,
        "home_ids": home_ids,
        "access_numbers": access_numbers,
        "target_table": target_table,
        "action_required": " ".join(action_lines) if action_lines else None,
        "status_to_delete": status_to_delete,
        "target_status": target_status,
        "source_status": source_status,
        "_code_blocks": code_blocks,
        "confidence": "HIGH",
    }


def _prepare_incident_context(incident_json: dict) -> tuple[dict, str | None]:
    """Merge a bare incident id with the matching report if one exists."""
    incident_id = _normalize_incident_id(incident_json.get("incident_id"))
    exact_report_path = _find_exact_report_path(incident_id)
    if not exact_report_path:
        return incident_json, None

    try:
        with open(exact_report_path, "r", encoding="utf-8") as f:
            report_data = json.load(f)
    except Exception:
        return incident_json, None

    report_context = _summarize_exact_report(report_data)
    merged = dict(report_context)
    for key, value in incident_json.items():
        if value not in (None, "", [], {}):
            merged[key] = value
    merged["incident_id"] = incident_id or incident_json.get("incident_id")
    return merged, exact_report_path


# ──────────────────────────────────────────────────────────────────────────────
# LAYER 3 — RETRIEVAL
# ──────────────────────────────────────────────────────────────────────────────

def retrieve(query: str, embed_model, embeddings, documents, metadata, top_k: int = TOP_K) -> list[dict]:
    q_emb = embed_model.encode([query], convert_to_numpy=True).astype("float32")[0]
    # Cosine similarity = dot product on normalized vectors
    norms = np.linalg.norm(embeddings, axis=1, keepdims=True) + 1e-9
    normed = embeddings / norms
    q_norm = q_emb / (np.linalg.norm(q_emb) + 1e-9)
    scores = normed @ q_norm
    indices = np.argsort(scores)[::-1][:top_k]
    return [
        {
            "text": documents[i],
            "source": metadata[i]["source"],
            "chunk_id": metadata[i]["chunk_id"],
            "score": float(scores[i]),
        }
        for i in indices
    ]


def build_retrieval_query(incident_json: dict) -> str:
    """Build a rich query string from the incident JSON for better retrieval."""
    parts = []
    if incident_json.get("incident_id"):
        parts.append(str(incident_json["incident_id"]))
    if incident_json.get("incident_type"):
        t = incident_json["incident_type"]
        parts.append(str(t) if isinstance(t, str) else " ".join(t))
    if incident_json.get("short_description"):
        parts.append(str(incident_json["short_description"]))
    if incident_json.get("action_required"):
        parts.append(incident_json["action_required"])
    if incident_json.get("full_description"):
        parts.append(incident_json["full_description"][:400])
    if incident_json.get("home_ids"):
        parts.append(" ".join(str(x) for x in incident_json["home_ids"]))
    if incident_json.get("access_numbers"):
        parts.append(" ".join(str(x) for x in incident_json["access_numbers"]))
    if incident_json.get("target_table"):
        parts.append(incident_json["target_table"])
    return " ".join(parts)


# ──────────────────────────────────────────────────────────────────────────────
# LAYER 4 — RESOLUTION (deterministic steps + Ollama narrative)
# ──────────────────────────────────────────────────────────────────────────────

RESOLUTION_SYSTEM_PROMPT = """You are an NRI incident resolution assistant.
You receive a structured incident JSON and retrieved knowledge chunks from past resolved incidents.
Return ONLY valid JSON — no markdown, no preamble.

Output schema:
{
  "incident_id": "string",
  "incident_type": "string",
  "confidence": "HIGH | MEDIUM | LOW",
  "resolution_method": "Method 1 | Method 2 | Custom | Unknown",
  "matched_sources": ["filenames that were useful"],
  "resolution_steps": [
    {
      "step": 1,
      "action": "plain English",
      "tool": "DBeaver | SQL | Excel | Function call | Manual",
      "sql_or_command": "exact SQL or null",
      "warning": "string or null"
    }
  ],
  "log_entry_template": "INSERT INTO support_tasks_log ... with <placeholders>",
  "coma_sync_required": true or false,
  "validation_checklist": ["item1", "item2"],
  "escalation_required": true or false,
  "escalation_reason": "string or null",
  "missing_information": ["list"],
  "notes": "string"
}

Resolution rules — apply these deterministically:

DELETE_PROVISION / DELETE_DEFECTIVE_PORT:
    table: use the report's target table, action: use the SQL shown in the report
    log: use the report title and incident id in the audit entry
  coma_sync_required: false

DELETE_DUPLICATE_RECORDS:
    table: use the report's target table, action: use the duplicate cleanup step shown in the report
    fallback if error: use the report's fallback step if present
    log: use the report title and incident id in the audit entry
  coma_sync_required: false

CHANGE_HOME_STATUS:
    Step 0: verify caller authorization if the report requires it
    Use the status-change step shown in the report
    coma_sync_required: true if the report mentions COMA verification

Use retrieved knowledge chunks for exact SQL, function signatures, and edge cases.
Never invent home IDs or table names not present in the inputs."""


def generate_resolution(incident_json: dict, chunks: list[dict], report_data: dict | None = None) -> dict:
    """Return a structured resolution derived from the matched report content."""
    report_driven = _build_report_resolution(incident_json, chunks, report_data=report_data)

    # If the incident type is already known from the report, prefer the report-driven
    # resolver so the output stays grounded in the exact report content instead of
    # letting the LLM invent method/escalation metadata.
    if report_driven.get("incident_type") != "UNKNOWN":
        return report_driven

    chunks_text = "\n\n".join(
        f"[CHUNK {i+1}]\nSource: {c['source']} (score: {c['score']:.2f})\n{c['text']}"
        for i, c in enumerate(chunks)
    )
    prompt = f"""{RESOLUTION_SYSTEM_PROMPT}

INCIDENT JSON:
{json.dumps(incident_json, indent=2)}

RETRIEVED KNOWLEDGE:
{chunks_text}
"""
    raw = ask_ollama(prompt)
    result = _parse_json_safe(raw)
    if "error" in result:
        return _build_report_resolution(incident_json, chunks, report_data)
    return result


def _infer_tool_from_code(code: str) -> str:
    text = code.strip().lower()
    if re.match(r"^select\s+public\.", text):
        return "Function call"
    if re.match(r"^(select|update|delete|insert)\b", text):
        return "SQL"
    return "Manual"


def _build_report_resolution(incident_json: dict, chunks: list[dict], report_data: dict | None = None) -> dict:
    """Build a report-grounded resolution without incident-specific hardcoding."""
    report_data = report_data or {}
    metadata = report_data.get("metadata", {}) if isinstance(report_data, dict) else {}
    blocks = report_data.get("blocks", []) if isinstance(report_data, dict) else []
    title = metadata.get("title") or incident_json.get("short_description") or incident_json.get("incident_id") or "Incident"

    incident_type = incident_json.get("incident_type") or metadata.get("incident_type") or "UNKNOWN"
    if isinstance(incident_type, list):
        incident_type = incident_type[0] if incident_type else "UNKNOWN"

    home_ids = list(incident_json.get("home_ids") or [])
    access_numbers = list(incident_json.get("access_numbers") or [])
    steps: list[dict] = []
    checklist: list[str] = []
    log_template = f"INSERT INTO support_tasks_log with {incident_json.get('incident_id', '<INC#>')}: {title}"
    source_status = incident_json.get("source_status")
    target_status = incident_json.get("target_status")
    target_table = incident_json.get("target_table")
    escalate = False
    escalate_reason = None
    coma = False

    for block in blocks:
        block_type = block.get("type")
        if block_type == "code":
            code = str(block.get("content", "")).strip()
            if code:
                steps.append({
                    "step": len(steps) + 1,
                    "action": code,
                    "tool": _infer_tool_from_code(code),
                    "sql_or_command": code,
                    "warning": None,
                })
                if not target_table:
                    lowered = code.lower()
                    if "fm_opv" in lowered:
                        target_table = "fm_opv"
                    elif "of_networkcreation" in lowered:
                        target_table = "of_networkCreation"
                if "coma_status" in code.lower():
                    coma = True
        elif block_type == "list":
            items = [str(item).strip() for item in block.get("items", []) if str(item).strip()]
            if items:
                checklist.extend(items)
                for item in items:
                    steps.append({
                        "step": len(steps) + 1,
                        "action": item,
                        "tool": "Manual",
                        "sql_or_command": None,
                        "warning": None,
                    })
        elif block_type == "description_box":
            items = [str(item).strip() for item in block.get("items", []) if str(item).strip()]
            checklist.extend(items)
        elif block_type == "table" and not home_ids:
            rows = block.get("rows", []) or []
            for row in rows:
                if isinstance(row, list) and row and HOME_ID_PATTERN.search(str(row[0])):
                    home_ids.append(str(row[0]).strip())
                if isinstance(row, list) and len(row) > 1 and str(row[1]).isdigit():
                    access_numbers.append(str(row[1]).strip())

    report_text = " ".join(
        str(block.get("content", ""))
        for block in blocks
        if block.get("type") in {"heading", "paragraph", "code"}
    ).lower()
    if incident_type == "UNKNOWN":
        if "duplicate" in report_text:
            incident_type = "DELETE_DUPLICATE_RECORDS"
        elif "defective port" in report_text:
            incident_type = "DELETE_DEFECTIVE_PORT"
        elif "status" in report_text and ("home" in report_text or "planning" in report_text or "unplanned" in report_text):
            incident_type = "CHANGE_HOME_STATUS"
        elif "provision" in report_text:
            incident_type = "DELETE_PROVISION"

    tools = {step["tool"] for step in steps if step.get("tool")}
    if tools == {"Function call"}:
        method = "Function call"
    elif tools == {"SQL"}:
        method = "SQL"
    elif tools == {"Manual"}:
        method = "Manual"
    elif tools:
        method = "Custom"
    else:
        method = "Unknown"

    if not checklist:
        checklist = [
            "Verify source data from the matched report",
            "Confirm the generated steps match the report's code blocks",
            "Validate affected records before applying changes",
        ]

    return {
        "incident_id": incident_json.get("incident_id") or metadata.get("incident_id") or "<INC#>",
        "incident_type": incident_type,
        "confidence": incident_json.get("confidence", metadata.get("confidence", "HIGH")),
        "resolution_method": method,
        "matched_sources": [c["source"] for c in chunks],
        "resolution_steps": steps,
        "log_entry_template": log_template,
        "coma_sync_required": coma,
        "validation_checklist": checklist,
        "escalation_required": escalate,
        "escalation_reason": escalate_reason,
        "missing_information": [],
        "notes": f"{len(chunks)} knowledge chunks retrieved. Report-driven resolution applied.",
        "target_table": target_table,
        "home_ids": home_ids,
        "access_numbers": access_numbers,
        "source_status": source_status,
        "target_status": target_status,
    }


# ──────────────────────────────────────────────────────────────────────────────
# OLLAMA WRAPPER
# ──────────────────────────────────────────────────────────────────────────────

def ask_ollama(prompt: str, model: str = OLLAMA_MODEL) -> str:
    try:
        import ollama
        resp = ollama.chat(model=model, messages=[{"role": "user", "content": prompt}])
        return resp["message"]["content"]
    except Exception as e:
        return f"Ollama error: {e}"


# ──────────────────────────────────────────────────────────────────────────────
# VLM (Qwen2-VL via MLX)
# ──────────────────────────────────────────────────────────────────────────────

@st.cache_resource(show_spinner=False)
def load_qwen_vlm(model_path: str = QWEN_MLX_MODEL_PATH):
    from mlx_vlm import load
    resolved = os.path.expanduser(model_path)
    if not os.path.isdir(resolved):
        raise ValueError(f"MLX model not found: {resolved}")
    model, processor = load(resolved)
    return model, processor


def run_qwen_mlx(prompt: str, image_path: str | None = None, max_tokens: int = 256) -> str:
    from mlx_vlm import generate, apply_chat_template
    model, processor = load_qwen_vlm()

    if getattr(processor, "chat_template", None) is None and hasattr(processor, "tokenizer"):
        processor.chat_template = getattr(processor.tokenizer, "chat_template", None)

    if image_path:
        messages = [{"role": "user", "content": [
            {"type": "image", "image": image_path},
            {"type": "text",  "text": prompt},
        ]}]
        prompt_text = apply_chat_template(processor, model.config, messages, add_generation_prompt=True, num_images=1)
        result = generate(model, processor, prompt_text, image=[image_path], verbose=False, max_tokens=max_tokens)
    else:
        prompt_text = apply_chat_template(processor, model.config, prompt, add_generation_prompt=True, num_images=0)
        result = generate(model, processor, prompt_text, verbose=False, max_tokens=max_tokens)

    return getattr(result, "text", str(result)).strip()


def _tmp_save(img: Image.Image) -> str:
    f = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
    img.save(f, format="PNG")
    return f.name


def _tmp_write(uploaded) -> str:
    suffix = os.path.splitext(uploaded.name or ".png")[1] or ".png"
    f = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
    f.write(uploaded.getvalue())
    f.flush()
    return f.name


# ──────────────────────────────────────────────────────────────────────────────
# UI HELPERS
# ──────────────────────────────────────────────────────────────────────────────

def _render_resolution(resolution: dict):
    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Incident type", str(resolution.get("incident_type", "—")))
    c2.metric("Confidence",    resolution.get("confidence", "—"))
    c3.metric("Method",        resolution.get("resolution_method", "—"))
    c4.metric("COMA sync",     "Yes" if resolution.get("coma_sync_required") else "No")

    if resolution.get("escalation_required"):
        st.error(f"🚨 Escalation required: {resolution.get('escalation_reason')}")

    if resolution.get("missing_information"):
        st.warning("⚠️ Missing information:\n" + "\n".join(f"- {x}" for x in resolution["missing_information"]))

    st.subheader("Resolution steps")
    for step in resolution.get("resolution_steps", []):
        label = f"Step {step['step']}: {step['action'][:70]}{'…' if len(step['action']) > 70 else ''}"
        with st.expander(label):
            st.write(f"**Tool:** {step['tool']}")
            if step.get("sql_or_command"):
                st.code(step["sql_or_command"], language="sql")
            if step.get("warning"):
                st.warning(f"⚠️ {step['warning']}")

    st.subheader("Validation checklist")
    for item in resolution.get("validation_checklist", []):
        st.checkbox(item, key=f"chk_{item[:40]}")

    if resolution.get("log_entry_template"):
        st.subheader("Log entry SQL")
        st.code(resolution["log_entry_template"], language="sql")

    if resolution.get("notes"):
        st.caption(resolution["notes"])

    st.subheader("Full JSON")
    st.json(resolution)


# ──────────────────────────────────────────────────────────────────────────────
# MAIN
# ──────────────────────────────────────────────────────────────────────────────

def main():
    st.set_page_config(page_title="NRI Incident Chatbot", layout="wide")
    st.title("NRI Incident Analysis & Resolution")
    st.caption("Main demo: Extract → Retrieve → Resolve. Qwen2-VL Test is only for direct vision checks.")

    with st.sidebar:
        st.subheader("How to use")
        st.write("1. Open **Main Demo**.")
        st.write("2. Paste incident text or upload a screenshot.")
        st.write("3. Review extracted JSON, retrieval, and resolution.")
        st.write("4. Use **Knowledge Base** to inspect indexed files.")
        st.write("5. Use **Qwen2-VL Test** only for raw model debugging.")

    # Load knowledge base once
    try:
        embed_model, embeddings, documents, metadata, n_files = build_knowledge_base()
        st.sidebar.success(f"Knowledge base: {n_files} files, {len(documents)} chunks")
    except Exception as e:
        st.sidebar.error(f"Knowledge base error: {e}")
        embed_model = embeddings = documents = metadata = None

    main_tab, kb_tab, qwen_tab = st.tabs([
        "Main Demo",
        "Knowledge Base",
        "Qwen2-VL Test",
    ])

    # ── TAB 1: MAIN PIPELINE ─────────────────────────────────────────────────
    with main_tab:
        st.info(
            "This is the actual demo flow. Paste a ticket description **or** upload a screenshot, and the app will extract fields, search the knowledge base, and generate a resolution automatically."
        )

        input_mode = st.radio("Input mode", ["Text", "Screenshot"], horizontal=True)
        incident_json = None
        temp_path = None
        # Step 1 — collect input
        if input_mode == "Text":
            user_text = st.text_area(
                "Paste incident description",
                height=180,
                placeholder="[GUXHG] Please delete the provisions in NRI from the attached file (Sheet Provisions) → Delete the P in column Status…",
            )
            extract_btn = st.button("Extract & Analyze", type="primary")

            if extract_btn and user_text.strip():
                with st.spinner("Extracting incident fields via Ollama…"):
                    incident_json = extract_from_text_via_ollama(user_text)
                if "error" not in incident_json:
                    st.subheader("Extracted incident JSON")
                    st.json(incident_json)
                else:
                    st.error(f"Extraction failed: {incident_json.get('raw','')}")

        else:  # Screenshot
            uploaded = st.file_uploader("Upload ServiceNow screenshot", type=["png", "jpg", "jpeg", "webp"])
            extract_btn = st.button("Extract from screenshot", type="primary", disabled=uploaded is None)

            if uploaded:
                img = Image.open(uploaded).convert("RGB")
                st.image(img, caption="Uploaded screenshot", use_column_width=True)

            if extract_btn and uploaded:
                temp_path = _tmp_write(uploaded)
                with st.spinner("Running Qwen2-VL vision extraction…"):
                    try:
                        raw_vlm = run_qwen_mlx(VLM_EXTRACTION_PROMPT, image_path=temp_path, max_tokens=512)
                        incident_json = _parse_json_safe(raw_vlm)
                        with st.expander("Raw VLM output"):
                            st.text(raw_vlm)
                        if "error" not in incident_json:
                            st.subheader("Extracted incident JSON")
                            st.json(incident_json)
                        else:
                            st.error("VLM output could not be parsed as JSON.")
                    except Exception as e:
                        st.error(f"VLM failed: {e} — Is the MLX model at {QWEN_MLX_MODEL_PATH}?")
                    finally:
                        if temp_path and os.path.exists(temp_path):
                            os.unlink(temp_path)

        # Step 2 — if we have extracted JSON, retrieve + resolve automatically
        if incident_json and "error" not in incident_json and embed_model is not None:
            st.divider()
            st.subheader("Knowledge base retrieval")

            incident_context, exact_report_path = _prepare_incident_context(incident_json)
            exact_report_data = None
            if exact_report_path:
                st.info(f"Exact report matched: {exact_report_path}")
                try:
                    with open(exact_report_path, "r", encoding="utf-8") as f:
                        exact_report_data = json.load(f)
                except Exception as e:
                    st.warning(f"Matched report could not be loaded: {e}")
                st.subheader("Resolved incident context")
                st.json(incident_context)

            query = build_retrieval_query(incident_context)
            chunks = retrieve(query, embed_model, embeddings, documents, metadata, top_k=TOP_K)

            exact_source_label = os.path.relpath(exact_report_path) if exact_report_path else None
            if exact_source_label:
                exact_chunks = [
                    {
                        "text": documents[i],
                        "source": metadata[i]["source"],
                        "chunk_id": metadata[i]["chunk_id"],
                        "score": 1.0,
                    }
                    for i in range(len(documents))
                    if metadata[i]["source"] == exact_source_label
                ]
                semantic_chunks = [c for c in chunks if c["source"] != exact_source_label]
                chunks = (exact_chunks + semantic_chunks)[:TOP_K]

            with st.expander(f"Retrieved {len(chunks)} chunks (click to inspect)"):
                for c in chunks:
                    # st.markdown(f"**{c['source']}** — chunk {c['chunk_id']} — score `{c['score']:.3f}`")
                    render_guided_steps(f"**{c['source']}** — chunk {c['chunk_id']} — score `{c['score']:.3f}`", c["text"])
                    st.text(c["text"][:300] + ("…" if len(c["text"]) > 300 else ""))
                    st.divider()

            st.divider()
            st.subheader("Resolution recommendation")

            with st.spinner("Generating resolution via Ollama…"):
                resolution = generate_resolution(incident_context, chunks, report_data=exact_report_data)

            _render_resolution(resolution)

        elif embed_model is None and incident_json:
            st.warning("Knowledge base unavailable — cannot retrieve context. Fix docs/ and reports/ folders and restart.")

    # ── TAB 2: KNOWLEDGE BASE INSPECTOR ──────────────────────────────────────
    with kb_tab:
        if embed_model is None:
            st.error("Knowledge base failed to load. Check docs/ and reports/ folders.")
        else:
            st.metric("Total chunks indexed", len(documents))
            st.caption("This index includes supported files from docs/ and reports/.")

            # File breakdown
            from collections import Counter
            counts = Counter(m["source"] for m in metadata)
            st.subheader("Files indexed")
            for fname, count in sorted(counts.items()):
                st.write(f"- **{fname}** — {count} chunks")

            st.divider()
            st.subheader("Search the knowledge base")
            kb_query = st.text_input("Query", placeholder="delete provision fm_opv")
            if kb_query:
                results = retrieve(kb_query, embed_model, embeddings, documents, metadata, top_k=5)
                for r in results:
                    with st.expander(f"{r['source']} — chunk {r['chunk_id']} — score {r['score']:.3f}"):
                        st.text(r["text"])

    # ── TAB 3: QWEN2-VL SANDBOX ──────────────────────────────────────────────
    with qwen_tab:
        st.info("Raw Qwen2-VL test bench — for debugging the vision model directly.")
        model_path_input = st.text_input("MLX model path", value=QWEN_MLX_MODEL_PATH)
        vlm_prompt = st.text_area(
            "Prompt",
            value="Inspect this screenshot. List every visible field label and value. Do not invent anything not visible.",
            height=120,
        )
        max_tok = st.slider("Max new tokens", 32, 512, 256, 32)
        up_img = st.file_uploader("Image", type=["png", "jpg", "jpeg", "webp"])

        col1, col2 = st.columns(2)
        run_up    = col1.button("Run on uploaded image", disabled=up_img is None)
        run_smoke = col2.button("Run smoke test")

        img_path = None
        smoke_cleanup = False

        if up_img:
            preview = Image.open(up_img).convert("RGB")
            st.image(preview, use_column_width=True)
            img_path = _tmp_write(up_img)
        elif run_smoke:
            smoke = Image.new("RGB", (640, 384), "#f4efe6")
            d = ImageDraw.Draw(smoke)
            d.rounded_rectangle((40, 40, 600, 300), radius=28, fill="#c8d6e5", outline="#2c3e50", width=4)
            d.ellipse((80, 90, 200, 210), fill="#ffcc80", outline="#8d6e63", width=4)
            d.text((86, 250), "MLX-VLM smoke test", fill="#1f2937")
            img_path = _tmp_save(smoke)
            st.image(smoke, caption="Smoke test image", use_column_width=True)
            smoke_cleanup = True

        if (run_up or run_smoke) and img_path:
            try:
                with st.spinner("Running Qwen2-VL…"):
                    out = run_qwen_mlx(vlm_prompt, image_path=img_path, max_tokens=max_tok)
                st.subheader("VLM output")
                st.write(out)
            except Exception as e:
                st.error(f"VLM error: {e}")
            finally:
                if img_path and os.path.exists(img_path):
                    os.unlink(img_path)


if __name__ == "__main__":
    main()