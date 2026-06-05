import os
import re
from PIL import Image, ImageDraw

import streamlit as st
from docx import Document
import numpy as np

# CONFIG
DOCS_FOLDER = "docs"
EMBED_MODEL_NAME = "all-MiniLM-L6-v2"
TOP_K = 3
QWEN_VL_MODEL_NAME = "Qwen/Qwen2.5-VL-3B-Instruct"

# Reduce noisy tokenizer threading warnings in constrained local environments.
os.environ.setdefault("TOKENIZERS_PARALLELISM", "false")

# LOAD DOCX
def extract_docx_content(file_path: str) -> str:
    """
    Extract paragraphs + tables from a .docx file.
    Screenshots/images OCR is skipped here for simplicity.
    """
    doc = Document(file_path)
    parts = []

    # Paragraphs
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            parts.append(text)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    return "\n".join(parts)

# CHUNKING
def chunk_text(text: str, chunk_size: int = 700, overlap: int = 120):
    """
    Basic character-based chunking.
    """
    text = re.sub(r"\s+", " ", text).strip()
    chunks = []

    start = 0
    while start < len(text):
        end = start + chunk_size
        chunk = text[start:end]
        chunks.append(chunk)
        start += (chunk_size - overlap)

    return chunks

# BUILD VECTOR INDEX
@st.cache_resource
def build_knowledge_base():
    from sentence_transformers import SentenceTransformer

    documents = []
    metadata = []

    if not os.path.isdir(DOCS_FOLDER):
        raise ValueError(f"Missing docs folder: {DOCS_FOLDER}")

    for filename in os.listdir(DOCS_FOLDER):
        if filename.endswith(".docx"):
            path = os.path.join(DOCS_FOLDER, filename)
            content = extract_docx_content(path)
            chunks = chunk_text(content)

            for i, chunk in enumerate(chunks):
                documents.append(chunk)
                metadata.append({
                    "source": filename,
                    "chunk_id": i
                })

    if not documents:
        raise ValueError("No .docx documents found in docs/ folder.")

    model = SentenceTransformer(EMBED_MODEL_NAME)
    embeddings = model.encode(documents, convert_to_numpy=True).astype("float32")

    return model, embeddings, documents, metadata

# RETRIEVAL
def retrieve(query: str, model, embeddings, documents, metadata, top_k: int = TOP_K):
    
    query_emb = model.encode([query], convert_to_numpy=True).astype("float32")[0]
    distances = np.linalg.norm(embeddings - query_emb, axis=1)
    indices = np.argsort(distances)[:top_k]

    results = []
    for idx in indices:
        results.append({
            "text": documents[idx],
            "source": metadata[idx]["source"],
            "chunk_id": metadata[idx]["chunk_id"],
            "distance": float(distances[idx])
        })
    return results

# PROMPT BUILDING
def build_prompt(user_query: str, retrieved_chunks):
    context = "\n\n".join(
        [f"[Source: {c['source']} | Chunk {c['chunk_id']}]\n{c['text']}" for c in retrieved_chunks]
    )

    prompt = f"""
You are an internal incident-support assistant.

Your role:
- Help the user understand possible resolution steps for a ServiceNow incident
- Use ONLY the provided context
- If unsure, say "Not enough information in knowledge base"
- Be practical and concise
- Output:
  1. Incident type guess
  2. Suggested resolution steps
  3. SQL/action examples if found
  4. Source references used

Context:
{context}

User incident:
{user_query}
"""
    return prompt.strip()

# LOCAL LLM CALL (OLLAMA)
def ask_ollama(prompt: str, model_name: str = "llama3:8b"):
    """
    Requires local Ollama running:
    - ollama serve
    - ollama pull llama3
    """
    try:
        import ollama
        response = ollama.chat(
            model=model_name,
            messages=[{"role": "user", "content": prompt}]
        )
        return response["message"]["content"]
    except Exception as e:
        return f"Local LLM error: {e}"


# QWEN2-VL TEST MODE
def _qwen_device():
    import torch

    if torch.cuda.is_available():
        return "cuda"
    if torch.backends.mps.is_available():
        return "mps"
    return "cpu"


@st.cache_resource(show_spinner=False)
def load_qwen2_vl(model_name: str = QWEN_VL_MODEL_NAME):
    import torch
    from transformers import AutoProcessor, Qwen2_5_VLForConditionalGeneration

    device = _qwen_device()
    # dtype = torch.float16 if device in {"cuda", "mps"} else torch.float32
    if device == "mps":
        dtype = torch.bfloat16
    elif device == "cuda":
        dtype = torch.float16
    else:
        dtype = torch.float32
        
    model = Qwen2_5_VLForConditionalGeneration.from_pretrained(
        model_name,
        torch_dtype=dtype,
        low_cpu_mem_usage=True,
        # device_map="auto" if device == "cpu" else None,
        device_map= None,
        trust_remote_code=True,
    )
    # if device != "cpu":
    model.to(device)
    processor = AutoProcessor.from_pretrained(model_name, trust_remote_code=True)
    return model, processor, device


def create_qwen_smoke_image() -> Image.Image:
    image = Image.new("RGB", (640, 384), color="#f4efe6")
    draw = ImageDraw.Draw(image)
    draw.rounded_rectangle((40, 40, 600, 300), radius=28, fill="#c8d6e5", outline="#2c3e50", width=4)
    draw.ellipse((80, 90, 200, 210), fill="#ffcc80", outline="#8d6e63", width=4)
    draw.rectangle((260, 110, 560, 230), fill="#81c784", outline="#2e7d32", width=4)
    draw.text((86, 250), "Qwen2-VL smoke test", fill="#1f2937")
    draw.text((290, 145), "Describe this scene", fill="#0f172a")
    return image


def run_qwen2_vl(image: Image.Image, prompt: str, model_name: str = QWEN_VL_MODEL_NAME, max_new_tokens: int = 128):
    import torch

    model, processor, device = load_qwen2_vl(model_name)
    image = image.convert("RGB")

    messages = [
        {
            "role": "user",
            "content": [
                # {"type": "image", "image": image},
                # {"type": "text", "text": prompt},
                {
                   "type": "image",
                   "image": image,
                   "min_pixels": 224 * 224,   # don't let it downsample too aggressively
                   "max_pixels": 1280 * 784,  # cap so it fits in 16 GB
                },
                {"type": "text", "text": prompt}
            ],
        }
    ]

    chat_text = processor.apply_chat_template(messages, tokenize=False, add_generation_prompt=True)
    inputs = processor(text=[chat_text], images=[image], return_tensors="pt")
    inputs = {key: value.to(device) for key, value in inputs.items()}

    with torch.inference_mode():
        generated_ids = model.generate(**inputs, max_new_tokens=max_new_tokens)

    prompt_length = inputs["input_ids"].shape[1]
    generated_text = processor.batch_decode(generated_ids[:, prompt_length:], skip_special_tokens=True)[0]
    return generated_text.strip()

# UI
def main():
    st.set_page_config(page_title="Incident Chatbot", layout="wide")
    st.title("ServiceNow Incident Analysis Chatbot")
    st.caption("Local RAG prototype plus a Qwen2-VL test bench")

    rag_tab, qwen_tab = st.tabs(["Incident RAG", "Qwen2-VL Test"])

    with rag_tab:
        st.info("Knowledge is loaded from .docx files in the docs folder.")

        try:
            model, embeddings, documents, metadata = build_knowledge_base()
        except Exception as e:
            st.error(f"Failed to build knowledge base: {e}")
            st.stop()

        user_query = st.text_area(
            "Paste incident description:",
            height=180,
            placeholder="Example: Please delete duplicate yellow-marked records from NRI..."
        )

        if st.button("Analyze Incident"):
            if not user_query.strip():
                st.warning("Please enter an incident description.")
                return

            retrieved = retrieve(user_query, model, embeddings, documents, metadata, top_k=TOP_K)
            prompt = build_prompt(user_query, retrieved)
            answer = ask_ollama(prompt, model_name="llama3:8b")

            st.subheader("Recommended Answer")
            st.write(answer)

            st.subheader("Retrieved Knowledge")
            for r in retrieved:
                with st.expander(f"{r['source']} | chunk {r['chunk_id']}"):
                    st.write(r["text"])

    with qwen_tab:
        st.info("Load Qwen/Qwen2.5-VL-3B-Instruct on demand and test it with an uploaded image or a built-in smoke test.")
        model_name = st.text_input("Qwen2-VL model id", value=QWEN_VL_MODEL_NAME)
        prompt = st.text_area(
            "Image prompt",
            value="Describe the image in one paragraph, then list any visible text or notable details.",
            height=140,
        )
        uploaded_image = st.file_uploader("Upload an image for Qwen2-VL", type=["png", "jpg", "jpeg", "webp"])

        col_left, col_right = st.columns(2)
        with col_left:
            run_uploaded = st.button("Run on uploaded image", disabled=uploaded_image is None)
        with col_right:
            run_smoke = st.button("Run smoke test on sample image")

        test_image = None
        if uploaded_image is not None:
            test_image = Image.open(uploaded_image)
            st.image(test_image, caption="Uploaded image", width="stretch")
        elif run_smoke:
            test_image = create_qwen_smoke_image()
            st.image(test_image, caption="Generated smoke test image", width="stretch")

        should_run = (run_uploaded or run_smoke) and test_image is not None
        if should_run:
            try:
                with st.spinner(f"Loading {model_name} and running inference..."):
                    result = run_qwen2_vl(test_image, prompt, model_name=model_name)
                st.subheader("Qwen2-VL Output")
                st.write(result)
            except Exception as e:
                st.error(f"Qwen2-VL test failed: {e}")

if __name__ == "__main__":
    main()